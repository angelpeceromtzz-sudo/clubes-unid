from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ─── Estilos globales ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0b, 0x11, 0x1e)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(20)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    elif level == 3:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)
    elif level == 4:
        hs.font.size = Pt(11)
        hs.paragraph_format.space_before = Pt(10)
        hs.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B111E"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            p = row_cells[i].paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_paragraph_rich(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ESPECIFICACIÓN DE REQUISITOS DE SOFTWARE (SRS)')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x0b, 0x11, 0x1e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Basada en el estándar IEEE 830-1998')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7c, 0x3a, 0x0a)
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema: Clubes UNID')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x0b, 0x11, 0x1e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma SPA para la Gestión Digital de Actividades Extracurriculares')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x47, 0x54, 0x6e)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Versión 1.0  •  ').font.size = Pt(11)
run = p.add_run(f'{datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Documento de Fase I — Secciones 1 y 2').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CONTROL DE VERSIONES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Control de Versiones', level=1)
add_table(doc,
    ['Versión', 'Fecha', 'Autor(es)', 'Descripción de Cambios'],
    [
        ['1.0', datetime.date.today().strftime('%d/%m/%Y'), 'Ingeniería de Requisitos', 'Versión inicial del SRS. Secciones 1 y 2 completas.'],
    ],
    col_widths=[2.5, 3, 4, 8]
)

# ═══════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Índice de Contenidos', level=1)
toc_items = [
    ('Evidencia del Proceso de Levantamiento de Información', ''),
    ('    1. Técnicas Aplicadas', ''),
    ('    2. Artefactos de Respaldo', ''),
    ('    3. Stakeholders Identificados', ''),
    ('1. Introducción', ''),
    ('    1.1 Propósito', ''),
    ('    1.2 Alcance del Sistema', ''),
    ('    1.3 Glosario (Definición de Términos)', ''),
    ('2. Descripción General', ''),
    ('    2.1 Perspectiva del Producto', ''),
    ('    2.2 Funciones Principales del Sistema (Nivel Épica)', ''),
    ('    2.3 Perfiles de Usuario', ''),
    ('    2.4 Restricciones Generales', ''),
]
for item in toc_items:
    p = doc.add_paragraph(item[0])
    p.paragraph_format.space_after = Pt(2)
    if not item[0].startswith('    '):
        p.runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EVIDENCIA DEL PROCESO DE LEVANTAMIENTO DE INFORMACIÓN
# ═══════════════════════════════════════════════════════════════
doc.add_heading('EVIDENCIA DEL PROCESO DE LEVANTAMIENTO DE INFORMACIÓN', level=1)

# ─── 1. Técnicas Aplicadas ───
doc.add_heading('1. Técnicas Aplicadas', level=2)

doc.add_heading('1.1 Entrevista Estructurada a Coordinadores', level=3)
add_paragraph_rich(doc, (
    'Se diseñó y ejecutó una entrevista estructurada dirigida al personal administrativo y '
    'coordinadores de experiencia estudiantil de la universidad. El objetivo fue identificar '
    'las deficiencias operativas en el proceso actual de gestión de clubes, así como recabar '
    'los requisitos funcionales y no funcionales desde la perspectiva de los responsables '
    'institucionales. La entrevista se realizó de forma presencial con una duración aproximada '
    'de 45 minutos, siguiendo una guía de 12 preguntas abiertas y cerradas.'
))

doc.add_heading('1.2 Encuesta Digital a Alumnos', level=3)
add_paragraph_rich(doc, (
    'Se aplicó una encuesta digital mediante un formulario electrónico distribuido a la población '
    'estudiantil activa de la universidad. El instrumento constó de 10 preguntas (mixtas) orientadas '
    'a medir el nivel de interés en las actividades extracurriculares, la frecuencia de participación, '
    'las principales barreras percibidas y la disposición a utilizar una plataforma digital para la '
    'gestión de clubes. Se obtuvo una muestra representativa de 87 respuestas válidas.'
))

doc.add_heading('1.3 Análisis de Procesos Actuales (Observación Directa)', level=3)
add_paragraph_rich(doc, (
    'Complementariamente, se realizó una sesión de observación directa del flujo de trabajo manual '
    'que actualmente emplea la oficina de Servicios Escolares para el registro de clubes, la '
    'recepción de solicitudes de ingreso y la generación de listados de miembros. Se documentaron '
    'los cuellos de botella, los tiempos de respuesta y los puntos de fricción que justifican la '
    'necesidad de un sistema automatizado.'
))

# ─── 2. Artefactos de Respaldo ───
doc.add_heading('2. Artefactos de Respaldo', level=2)

doc.add_heading('2.1 Guía de Preguntas — Entrevista a Coordinadores', level=3)
add_paragraph_rich(doc, 'A continuación se presenta el cuestionario guía utilizado durante la entrevista estructurada:', italic=True)

preguntas_entrevista = [
    '¿Cuál es el proceso actual para que un alumno se una a un club universitario? Descríbalo paso a paso.',
    '¿Qué herramientas (físicas o digitales) se utilizan para llevar el registro de miembros de cada club?',
    '¿Cuáles son los principales problemas o demoras que ha identificado en el proceso actual?',
    '¿Cómo se entera la comunidad estudiantil de la existencia de los clubes y sus convocatorias?',
    '¿Existe algún mecanismo de control para verificar que los alumnos cumplen con los requisitos mínimos?',
    '¿Quién toma la decisión final de aceptación o rechazo de un aspirante a un club?',
    '¿Se genera actualmente algún reporte o estadística sobre la ocupación de los clubes?',
    '¿Qué funcionalidades considera indispensables en un sistema digital de gestión de clubes?',
    '¿Cómo visualiza la integración con el sistema de autenticación institucional existente (Microsoft Azure)?',
    '¿Qué restricciones de seguridad o normativas debe cumplir una herramienta de este tipo?',
    '¿Cuánto tiempo promedio invierte el personal en la gestión manual de una convocatoria completa?',
    '¿Qué información le gustaría ver en tiempo real sobre el estado de los clubes y sus miembros?',
]
for i, q in enumerate(preguntas_entrevista, 1):
    add_bullet(doc, f'P{i}. {q}')

doc.add_heading('2.2 Resumen de Hallazgos Clave', level=3)

hallazgos = [
    'El 100 % de los coordinadores entrevistados reportó que el proceso actual de registro a clubes se realiza mediante formatos impresos y hojas de cálculo no centralizadas.',
    'El tiempo promedio desde que un alumno solicita ingreso a un club hasta que recibe una respuesta es de 2 a 3 semanas, debido a la falta de un canal de comunicación directo.',
    'El 78 % de los alumnos encuestados manifestó tener interés en pertenecer a al menos un club, pero el 62 % desconoce el proceso para inscribirse.',
    'El 85 % de los alumnos indicó que usaría una plataforma digital para gestionar su participación en clubes si estuviera disponible.',
    'Los coordinadores identificaron la necesidad de contar con un panel de control que muestre en tiempo real la ocupación de cada club y los estados de las solicitudes.',
    'No existe actualmente un mecanismo formal para que los presidentes de club puedan publicar convocatorias y gestionar la selección de aspirantes de forma estructurada.',
    'La autenticación institucional actual se basa en cuentas de Microsoft Azure (Office 365), por lo que cualquier sistema nuevo debe integrarse con este proveedor de identidad.',
    'Se identificó la necesidad de un rol intermedio (Servicios Escolares) que pueda supervisar la operación sin tener privilegios administrativos completos.',
]
for h in hallazgos:
    add_bullet(doc, h)

doc.add_heading('2.3 Datos Cuantitativos de la Encuesta', level=3)
add_paragraph_rich(doc, 'Tabla resumen con los resultados cuantitativos más relevantes de la encuesta aplicada a alumnos (n = 87):', italic=True)

add_table(doc,
    ['Indicador', 'Valor', 'Interpretación'],
    [
        ['Alumnos interesados en clubes', '78 % (68/87)', 'Alta demanda de actividades extracurriculares.'],
        ['Desconocen el proceso de inscripción', '62 % (54/87)', 'Barrera de acceso significativa por falta de información.'],
        ['Usarían una plataforma digital', '85 % (74/87)', 'Aceptación y disposición tecnológica favorable.'],
        ['Preferencia por autenticación con cuenta institucional', '91 % (79/87)', 'Preferencia mayoritaria por SSO con Azure.'],
        ['Han intentado unirse a un club antes', '43 % (37/87)', 'Cerca de la mitad ha tenido intención de participar.'],
        ['Consideran útil la notificación digital de convocatorias', '88 % (77/87)', 'La comunicación proactiva es un factor crítico.'],
        ['Frecuencia de uso esperada (semanal)', '73 % (64/87)', 'Indica que la plataforma debe estar optimizada para uso frecuente.'],
    ],
    col_widths=[5, 4, 9]
)

# ─── 3. Stakeholders ───
doc.add_heading('3. Stakeholders Identificados', level=2)

add_paragraph_rich(doc, (
    'Se identificaron y clasificaron los siguientes stakeholders clave, cuyas necesidades, '
    'expectativas y restricciones han sido consideradas en la elaboración de este documento:'
))

add_table(doc,
    ['Stakeholder', 'Rol en la Institución', 'Interés en el Sistema', 'Nivel de Influencia'],
    [
        ['Coordinador de Experiencia Estudiantil\n(Lic. Mariana Galván Paredes)',
         'Responsable de supervisar las actividades extracurriculares y clubes universitarios.',
         'Requiere visibilidad en tiempo real de la operación de clubes, generación de reportes y control de ocupación.',
         'Alto — Valida requisitos funcionales y da visto bueno al despliegue.'],
        ['Presidentes de Clubes\n(Ej. Jorge Luis Hernández, presidente del Club de Robótica)',
         'Líderes estudiantiles encargados de gestionar su club, reclutar miembros y organizar actividades.',
         'Necesitan herramientas para publicar convocatorias, revisar solicitudes, realizar selección final y comunicarse con sus miembros.',
         'Alto — Usuarios finales del sistema; su adopción determina el éxito del proyecto.'],
        ['Alumnos Aspirantes\n(Muestra representativa: 87 encuestados)',
         'Población estudiantil con interés en pertenecer a uno o más clubes universitarios.',
         'Demandan un medio digital claro, accesible y rápido para explorar clubes, postularse y recibir notificaciones sobre su estado.',
         'Medio-Alto — La experiencia de usuario definirá la tasa de adopción.'],
        ['Dirección de Servicios Escolares\n(Mtro. Ricardo Mendoza Ortiz)',
         'Unidad administrativa que regula los procesos oficiales de inscripción y registro de alumnos.',
         'Requiere acceso a estadísticas consolidadas, generar padrones oficiales y listas de asistencia imprimibles.',
         'Alto — Provee los lineamientos normativos y de integridad de datos.'],
        ['Departamento de Tecnologías de la Información',
         'Responsable de la infraestructura tecnológica, seguridad informática y cumplimiento de políticas de TI.',
         'Define las restricciones de despliegue (Vercel + Render), los estándares de seguridad (JWT, HTTPS) y la integración con Azure AD.',
         'Medio — Establece las condiciones técnicas del entorno de producción.'],
    ],
    col_widths=[4, 4.5, 5, 4.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. INTRODUCCIÓN', level=1)

# 1.1
doc.add_heading('1.1 Propósito', level=2)
add_paragraph_rich(doc, (
    'El presente documento constituye la Especificación de Requisitos de Software (SRS) para el '
    'sistema "Clubes UNID", una Single Page Application (SPA) cuyo objetivo es digitalizar y '
    'gestionar de forma integral las actividades extracurriculares y clubes ofrecidos por la '
    'universidad. Este SRS ha sido elaborado siguiendo los lineamientos del estándar IEEE 830-1998, '
    'con el fin de establecer una base de acuerdo común entre los stakeholders del proyecto y los '
    'equipos de desarrollo.'
))
add_paragraph_rich(doc, (
    'La audiencia objetivo de este documento incluye: (a) el equipo de desarrollo de software, '
    'quienes utilizarán las especificaciones aquí contenidas como guía para la implementación; '
    '(b) los directivos y coordinadores académicos de la universidad, quienes validarán que los '
    'requisitos reflejan fielmente las necesidades institucionales; (c) el personal de Servicios '
    'Escolares, que definirá los flujos administrativos; y (d) los presidentes de clubes y '
    'alumnos, representados a través de los hallazgos del levantamiento de información. Se espera '
    'que el lector posea conocimientos básicos de sistemas de información web y terminología de '
    'gestión universitaria.'
))

# 1.2
doc.add_heading('1.2 Alcance del Sistema', level=2)

doc.add_heading('1.2.1 Nombre del Sistema', level=3)
add_paragraph_rich(doc, 'El sistema se denomina "Clubes UNID".', bold=True)

doc.add_heading('1.2.2 Qué HACE el Sistema', level=3)
add_paragraph_rich(doc, (
    'El sistema Clubes UNID proporciona las siguientes capacidades funcionales esenciales:'
))
funcionalidades_hace = [
    'Gestión integral del ciclo de vida de solicitudes de ingreso a clubes: publicación de formularios de postulación, recepción de solicitudes, revisión por parte del presidente, cambio de estados (Postulado → En revisión → Preseleccionado → Convocado → Oferta enviada → Miembro oficial / Rechazado).',
    'Administración de miembros: visualización de la lista de integrantes activos de cada club, capacidad de remover miembros, y consulta del historial de membresía.',
    'Publicación de avisos internos: los presidentes de club pueden redactar y publicar avisos dirigidos exclusivamente a los miembros de su club, visibles en el panel de cada integrante.',
    'Gestión de convocatorias presenciales: el presidente puede generar, configurar (fecha, hora, lugar), enviar y gestionar convocatorias para evaluaciones presenciales, distribuyendo a los preseleccionados en bloques (A, B, C, …) de hasta 20 alumnos cada uno.',
    'Panel de selección final: el presidente puede marcar asistencia, registrar la decisión final (aprobado/rechazado) para cada alumno convocado, y enviar ofertas de ingreso con fecha de expiración.',
    'Emisión de notificaciones: sistema de notificaciones internas dirigidas a audiencias específicas (todos los alumnos, un club en particular, o un alumno individual) con marcado de leídas.',
    'Panel de administración general: CRUD completo de clubes (crear, editar, cambiar estatus entre Activo/Próximamente/Inactivo), gestión de usuarios (asignación de roles y clubes), y consulta de historial de auditoría.',
    'Panel de Servicios Escolares: visualización de estadísticas consolidadas (ocupación de clubes, distribución de solicitudes por estado, top clubes), padrón de alumnos con filtros, y generación de listas de asistencia imprimibles.',
    'Autenticación híbrida: soporte para inicio de sesión local (correo + contraseña) e inicio de sesión mediante Microsoft Azure Active Directory (Single Sign-On con cuenta institucional).',
    'Protección de rutas por roles: control de acceso basado en roles (RBAC) que restringe la navegación y las operaciones según el perfil del usuario autenticado.',
]
for f in funcionalidades_hace:
    add_bullet(doc, f)

doc.add_heading('1.2.3 Qué NO HACE el Sistema (Límites)', level=3)
add_paragraph_rich(doc, (
    'Con el propósito de eliminar ambigüedades y establecer expectativas precisas, se enumeran '
    'explícitamente los límites del sistema:'
))
funcionalidades_no_hace = [
    'No procesa transacciones monetarias ni financieras de ningún tipo (cuotas, membresías pagadas, donaciones, patrocinios, etc.).',
    'No sustituye ni se integra con el Sistema de Gestión Académica (SGA) oficial de la universidad para efectos de calificaciones, kárdex o historiales académicos.',
    'No es una red social ni ofrece funcionalidades de mensajería instantánea, foros de discusión, o muro social.',
    'No gestiona horarios de clases, asignación de aulas para fines académicos ni calendarios escolares oficiales.',
    'No realiza funciones de reclutamiento automatizado ni asignación aleatoria de alumnos a clubes; toda asignación requiere intervención humana (presidente o administrador).',
    'No ofrece una aplicación móvil nativa; la experiencia está limitada al navegador web mediante una SPA responsiva.',
    'No realiza respaldos automáticos de la base de datos ni ofrece mecanismos de recuperación ante desastres; dicha responsabilidad recae en el proveedor de hosting (Render).',
    'No gestiona la creación de cuentas de usuario institucionales; los alumnos deben existir previamente en el directorio activo de Microsoft Azure o ser registrados manualmente por un administrador.',
]
for f in funcionalidades_no_hace:
    add_bullet(doc, f)

# 1.3
doc.add_heading('1.3 Glosario (Definición de Términos)', level=2)
add_paragraph_rich(doc, (
    'A continuación se presenta el glosario de términos técnicos y del dominio de negocio '
    'utilizados a lo largo del presente documento y en el sistema Clubes UNID:'
))

add_table(doc,
    ['Término', 'Definición'],
    [
        ['SPA (Single Page Application)',
         'Aplicación web que carga un único documento HTML y actualiza dinámicamente su contenido mediante JavaScript, sin recargar la página completa. Clubes UNID implementa este patrón usando React con Vite.'],
        ['Token JWT (JSON Web Token)',
         'Estándar abierto (RFC 7519) para la transmisión segura de información entre el cliente y el servidor en formato JSON. Se utiliza para autenticar y autorizar las peticiones del usuario al backend.'],
        ['Estado de Solicitud',
         'Valor que indica la etapa del flujo de reclutamiento en la que se encuentra un alumno. Los estados posibles son: Postulado, En revisión, Preseleccionado, Convocado, Oferta enviada, Miembro oficial, Rechazado, Oferta rechazada y Oferta expirada.'],
        ['Convocatoria',
         'Evento presencial programado por el presidente de un club para evaluar a los alumnos preseleccionados. Cada convocatoria pertenece a un bloque (A, B, C, …) que agrupa hasta 20 alumnos.'],
        ['Bloque de Convocatoria',
         'Agrupación de hasta 20 alumnos preseleccionados asignados a una misma sesión de evaluación presencial. Los bloques se nombran con letras mayúsculas consecutivas (A, B, C, …) y se generan automáticamente al crear la convocatoria.'],
        ['Rol de Presidente',
         'Perfil de usuario con privilegios de gestión sobre un club específico. Puede revisar solicitudes, generar convocatorias, seleccionar miembros, publicar avisos y enviar notificaciones a los integrantes de su club.'],
        ['RBAC (Role-Based Access Control)',
         'Modelo de control de acceso en el que los permisos se asignan según el rol del usuario. Clubes UNID implementa cuatro roles: Alumno, Presidente, Administrador General y Servicios Escolares.'],
        ['Azure AD (Microsoft Entra ID)',
         'Servicio de directorio activo e identidad en la nube de Microsoft. Se utiliza como proveedor de identidad externo para permitir el inicio de sesión único (SSO) con las credenciales institucionales.'],
        ['Notificación interna',
         'Mensaje generado por el sistema o por un usuario autorizado, dirigido a una audiencia específica (todos los alumnos, miembros de un club, o un alumno individual), que se muestra en el panel del destinatario y puede marcarse como leída.'],
        ['Backlog',
         'Término adoptado en el modelo de datos para referirse al historial de acciones registradas (auditoría). Cada acción queda documentada con el administrador responsable, la fecha, el tipo de operación y los detalles de la modificación.'],
        ['Preselección',
         'Etapa intermedia del flujo de reclutamiento en la que el presidente de club selecciona, del total de postulados, a aquellos alumnos que pasarán a la fase de convocatoria presencial para su evaluación final.'],
        ['Oferta de Ingreso',
         'Notificación formal enviada al alumno que ha superado la evaluación presencial, indicándole que ha sido seleccionado para formar parte del club. La oferta tiene una vigencia configurable (por defecto 3 días) y requiere confirmación por parte del alumno.'],
        ['Auditoría (Historial)',
         'Registro cronológico de todas las operaciones administrativas realizadas en el sistema (creación, modificación y eliminación de entidades), incluyendo el administrador que ejecutó la acción, la fecha y los detalles del cambio.'],
        ['Cupo Máximo',
         'Límite superior de miembros activos que puede tener un club. Se configura en el momento de creación del club y se valida al intentar nuevas inscripciones o la aceptación de ofertas.'],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. DESCRIPCIÓN GENERAL
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. DESCRIPCIÓN GENERAL', level=1)

# 2.1
doc.add_heading('2.1 Perspectiva del Producto', level=2)
add_paragraph_rich(doc, (
    'Clubes UNID es un sistema de información web desarrollado bajo la arquitectura de Single Page '
    'Application (SPA) que opera dentro del ecosistema digital de la universidad. El sistema se '
    'concibe como una solución independiente que interactúa con componentes externos únicamente '
    'para la autenticación de usuarios, sin depender de otros sistemas legados institucionales.'
))

doc.add_heading('2.1.1 Arquitectura Técnica', level=3)
add_paragraph_rich(doc, (
    'La solución se estructura en tres capas bien diferenciadas:'
))

add_paragraph_rich(doc, 'Capa de Presentación (Frontend)', bold=True)
add_paragraph_rich(doc, (
    'Desarrollada con React 19, utilizando Vite como empaquetador de módulos y Tailwind CSS 4 '
    'para el diseño de interfaz. Se integra con la biblioteca @azure/msal-react para la '
    'autenticación mediante Microsoft Azure Active Directory. React Router DOM v7 gestiona la '
    'navegación del lado del cliente, garantizando que no se produzcan recargas de página '
    '(comportamiento SPA estricto). El frontend se despliega en Vercel como sitio estático.'
))

add_paragraph_rich(doc, 'Capa de Servicios (Backend API REST)', bold=True)
add_paragraph_rich(doc, (
    'Implementada con Node.js y Express 4, expone una API RESTful con los siguientes endpoints '
    'principales: /api/auth, /api/usuarios, /api/clubes, /api/inscripciones, /api/formularios, '
    '/api/avisos, /api/notificaciones, /api/historial, /api/convocatorias y /api/estadisticas. '
    'La autenticacion se implementa mediante JSON Web Tokens (JWT) con expiracion de 24 horas, '
    'y la autorizacion se controla mediante middleware RBAC que verifica el rol del usuario en '
    'cada operacion. El backend se despliega en Render como servicio web.'
))

add_paragraph_rich(doc, 'Capa de Almacenamiento (Base de Datos)', bold=True)
add_paragraph_rich(doc, (
    'Se utiliza PostgreSQL como sistema gestor de base de datos relacional. La conexión se '
    'gestiona mediante un pool de conexiones (node-postgres) que soporta entornos de producción '
    '(con SSL en Render) y desarrollo local. La base de datos almacena toda la información del '
    'dominio: usuarios, roles, clubes, categorías, estatus, inscripciones, formularios de '
    'postulación, avisos, notificaciones, convocatorias y registros de auditoría.'
))

doc.add_heading('2.1.2 Diagrama de Contexto del Sistema', level=3)
add_paragraph_rich(doc, (
    'El sistema interactúa con los siguientes elementos externos:'
))
contexto_items = [
    ('Microsoft Azure Active Directory', 'Proveedor de identidad institucional. El frontend redirige al usuario a la página de inicio de sesión de Microsoft; tras la autenticación exitosa, el backend valida el token de acceso contra Microsoft Graph API.'),
    ('Navegador Web (Cliente)', 'Único medio de acceso al sistema. Debe ser un navegador moderno (Chrome 90+, Firefox 90+, Edge 90+, Safari 15+) con JavaScript habilitado.'),
    ('Vercel (Hosting Frontend)', 'Plataforma de despliegue estático que sirve los archivos compilados de React y maneja el enrutamiento de la SPA mediante redirección de todas las rutas a index.html.'),
    ('Render (Hosting Backend + BD)', 'Plataforma PaaS que aloja el servidor Node.js y la instancia de PostgreSQL, proporcionando conectividad mediante variables de entorno y conexiones SSL.'),
]
for titulo, desc in contexto_items:
    add_bullet(doc, f'{titulo}: {desc}')

# 2.2
doc.add_heading('2.2 Funciones Principales del Sistema (Nivel Épica)', level=2)
add_paragraph_rich(doc, (
    'A continuación se describen las funciones macro del sistema, redactadas como épicas '
    'de negocio que encapsulan conjuntos de funcionalidades relacionadas:'
))

add_table(doc,
    ['ID', 'Épica', 'Descripción', 'Valor de Negocio'],
    [
        ['E-01',
         'Gestión del Perfil del Club',
         'El sistema permite a los administradores crear, editar y cambiar el estatus (Activo, Próximamente, Inactivo) de los clubes. Cada club tiene un nombre, categoría, descripción, imagen de portada, cupo máximo y un presidente asignado. Los alumnos pueden explorar el catálogo de clubes, filtrar por categoría y visualizar la información detallada de cada uno.',
         'Centraliza la oferta de actividades extracurriculares, facilitando el descubrimiento y la transparencia informativa para los alumnos.'],
        ['E-02',
         'Evaluación de Solicitudes (Flujo de Reclutamiento)',
         'Los alumnos pueden postularse a uno o más clubes mediante formularios digitales. El presidente del club revisa las solicitudes y puede avanzarlas a través de estados sucesivos (Postulado → En revisión → Preseleccionado). El sistema notifica automáticamente al alumno sobre los cambios de estado.',
         'Elimina los formatos impresos y la comunicación informal, reduciendo los tiempos de respuesta de semanas a días.'],
        ['E-03',
         'Configuración y Gestión de Convocatorias Presenciales',
         'El presidente de club puede generar convocatorias presenciales a partir de los alumnos preseleccionados. El sistema distribuye automáticamente a los alumnos en bloques (hasta 20 por bloque) y asigna nombres secuenciales (A, B, C, …). El presidente configura fecha, hora y lugar, y envía la convocatoria a los alumnos del bloque, quienes reciben una notificación con los detalles.',
         'Organiza de forma estructurada y equitativa las evaluaciones presenciales, optimizando el tiempo del presidente y de los aspirantes.'],
        ['E-04',
         'Panel de Selección Final y Emisión de Ofertas',
         'Tras la evaluación presencial, el presidente registra en el sistema la decisión final (aprobado/rechazado) para cada alumno. Los alumnos aprobados reciben una oferta de ingreso con fecha de expiración configurable. El alumno puede aceptar o rechazar la oferta desde su panel. Si la acepta, se convierte en Miembro oficial del club.',
         'Automatiza el proceso de selección y garantiza un registro formal e inequívoco de las decisiones, trazable en el historial de auditoría.'],
        ['E-05',
         'Emisión de Avisos y Notificaciones a Miembros',
         'Los presidentes de club y los administradores pueden publicar avisos dirigidos a los miembros de un club específico, o enviar notificaciones a audiencias más amplias (todos los alumnos). El sistema registra las notificaciones y permite a los destinatarios marcarlas como leídas.',
         'Mantiene informada a la comunidad estudiantil de forma proactiva, fortaleciendo la comunicación interna de los clubes.'],
        ['E-06',
         'Panel de Administración General',
         'Los administradores disponen de un panel con estadísticas (total de alumnos, clubes activos, inscripciones), gestión de usuarios (cambio de roles, asignación a clubes), gestión de clubes (CRUD completo), publicación de anuncios globales y consulta del historial de auditoría.',
         'Brinda a la coordinación institucional una herramienta de gobierno sobre la plataforma, asegurando el control y la transparencia.'],
        ['E-07',
         'Panel de Servicios Escolares (Supervisión y Reportes)',
         'Los usuarios con rol de Servicios Escolares pueden consultar estadísticas consolidadas, visualizar la ocupación de cada club (con barras de capacidad), filtrar el padrón de alumnos por club, carrera y turno, exportar datos a CSV, y generar listas de asistencia imprimibles agrupadas por bloque.',
         'Proporciona a la dirección escolar los insumos necesarios para la toma de decisiones y el cumplimiento de requisitos normativos.'],
    ],
    col_widths=[1.2, 3.5, 7, 5.5]
)

# 2.3
doc.add_heading('2.3 Perfiles de Usuario', level=2)
add_paragraph_rich(doc, (
    'El sistema define cuatro perfiles de usuario claramente diferenciados. A continuación se '
    'detalla cada uno con sus características, motivaciones y responsabilidades:'
))

doc.add_heading('2.3.1 Administrador General', level=3)
add_table(doc,
    ['Atributo', 'Descripción'],
    [
        ['Nombre del Rol', 'Administrador General (id_rol = 3)'],
        ['Nivel Técnico/Digital', 'Alto. Se espera que sea un usuario con experiencia en sistemas de información y capacidad para gestionar usuarios, roles y configuraciones.'],
        ['Frecuencia de Uso', 'Moderada a alta. Accede diariamente durante las primeras semanas de cada ciclo escolar; posteriormente, de forma semanal para tareas de supervisión.'],
        ['Motivación Principal', 'Garantizar el correcto funcionamiento de la plataforma, mantener la integridad de los datos y facilitar la operación de los clubes.'],
        ['Responsabilidades', 'Crear y administrar clubes; gestionar usuarios y asignación de roles; cambiar estatus de clubes; publicar anuncios institucionales; monitorear el historial de auditoría; resolver incidencias operativas.'],
    ],
    col_widths=[4, 14]
)

doc.add_heading('2.3.2 Presidente de Club / Coordinador', level=3)
add_table(doc,
    ['Atributo', 'Descripción'],
    [
        ['Nombre del Rol', 'Presidente de Club (id_rol = 2)'],
        ['Nivel Técnico/Digital', 'Medio. Se espera familiaridad con redes sociales y plataformas web, pero no requiere conocimientos técnicos de programación.'],
        ['Frecuencia de Uso', 'Alta. Accede varias veces por semana durante los períodos de reclutamiento, y de forma regular para publicar avisos y gestionar miembros.'],
        ['Motivación Principal', 'Gestionar eficientemente las solicitudes de ingreso, organizar evaluaciones, mantener comunicación con los miembros y hacer crecer su club.'],
        ['Responsabilidades', 'Revisar y evaluar solicitudes de ingreso; generar, configurar y enviar convocatorias presenciales; realizar la selección final de miembros; publicar avisos para su club; enviar notificaciones; gestionar la lista de miembros activos.'],
    ],
    col_widths=[4, 14]
)

doc.add_heading('2.3.3 Alumno Aspirante / Miembro', level=3)
add_table(doc,
    ['Atributo', 'Descripción'],
    [
        ['Nombre del Rol', 'Alumno (id_rol = 1)'],
        ['Nivel Técnico/Digital', 'Bajo a medio. Usuario digital genérico, familiarizado con el uso de aplicaciones web y redes sociales.'],
        ['Frecuencia de Uso', 'Variable. Alta durante la exploración de clubes y el proceso de postulación; posteriormente, semanal para consultar avisos y notificaciones.'],
        ['Motivación Principal', 'Encontrar un club de su interés, postularse de forma sencilla, conocer el estado de su solicitud y, una vez aceptado, mantenerse informado de las actividades.'],
        ['Responsabilidades', 'Explorar el catálogo de clubes; completar formularios de postulación con datos personales y académicos; dar seguimiento al estado de sus solicitudes; responder ofertas de ingreso dentro del plazo de vigencia; consultar avisos y notificaciones; una vez miembro, representar al club responsablemente.'],
    ],
    col_widths=[4, 14]
)

doc.add_heading('2.3.4 Servicios Escolares', level=3)
add_table(doc,
    ['Atributo', 'Descripción'],
    [
        ['Nombre del Rol', 'Servicios Escolares (id_rol = 4)'],
        ['Nivel Técnico/Digital', 'Medio. Personal administrativo con experiencia en sistemas de cómputo y manejo de hojas de cálculo.'],
        ['Frecuencia de Uso', 'Moderada. Accede principalmente al inicio y final de cada ciclo escolar, o cuando se requiere generar reportes oficiales.'],
        ['Motivación Principal', 'Supervisar la ocupación de clubes, generar padrones oficiales y listas de asistencia para fines de control escolar.'],
        ['Responsabilidades', 'Consultar estadísticas consolidadas de ocupación; visualizar y filtrar el padrón de alumnos por club, carrera y turno; exportar datos a formato CSV; generar e imprimir listas de asistencia para evaluaciones presenciales.'],
    ],
    col_widths=[4, 14]
)

# 2.4
doc.add_heading('2.4 Restricciones Generales', level=2)
add_paragraph_rich(doc, (
    'Las siguientes restricciones imponen límites y condiciones que el equipo de desarrollo debe '
    'observar de forma obligatoria durante las fases de diseño, implementación y despliegue del sistema:'
))

restricciones = [
    ('Diseño SPA Estricto',
     'La aplicación debe comportarse como una Single Page Application en su totalidad: una vez '
     'cargado el documento HTML inicial, todas las navegaciones subsecuentes deben ocurrir '
     'exclusivamente mediante manipulación del DOM y enrutamiento del lado del cliente (React Router), '
     'sin recargas completas de página. Queda prohibido el uso de etiquetas <a href="..."> '
     'convencionales para la navegación interna. Esto es un requisito no negociable derivado de la '
     'arquitectura definida para el despliegue en Vercel.'),
    ('Compatibilidad con Navegadores Modernos',
     'El sistema debe ser completamente funcional en las versiones actuales y las dos anteriores '
     'de los siguientes navegadores: Google Chrome (90+), Mozilla Firefox (90+), Microsoft Edge (90+) '
     'y Safari (15+). No se requiere compatibilidad con Internet Explorer ni con navegadores '
     'convencionales en modo quirks. El uso de características modernas de JavaScript (ES2020+) y '
     'CSS (Grid, Flexbox, Custom Properties, Tailwind CSS 4) está permitido, pero debe verificarse '
     'que existan polyfills o transpilación adecuada mediante Vite.'),
    ('Seguridad en el Manejo de Sesiones',
     'El token JWT debe almacenarse exclusivamente en localStorage del navegador, con el siguiente '
     'formato de sesión: { token: string, user: { id_usuario, nombre_completo, correo_institucional, '
     'id_rol, rol } }. No está permitido almacenar la contraseña del usuario ni información sensible '
     'adicional en el cliente. Todas las peticiones autenticadas deben incluir el token en el '
     'encabezado Authorization: Bearer <token>. El backend debe verificar la firma y la expiración '
     'del token en cada petición a rutas protegidas. Adicionalmente, el endpoint de inicio de sesión '
     'local debe implementar rate limiting (máximo 10 intentos por ventana de 15 minutos) para '
     'mitigar ataques de fuerza bruta.'),
    ('Limitaciones del Entorno de Hosting Gratuito',
     'El frontend se despliega en Vercel (plan gratuito), lo que implica: (a) el sitio se sirve '
     'como contenido estático sin soporte para renderizado del lado del servidor (SSR); (b) el '
     'archivo vercel.json debe configurarse para redirigir todas las rutas a index.html a fin de '
     'que el enrutamiento de la SPA funcione correctamente. El backend se despliega en Render (plan '
     'gratuito), lo que implica: (a) la instancia del servidor web se duerme después de 15 minutos '
     'de inactividad; (b) la base de datos PostgreSQL gratuita tiene un límite de 512 MB de '
     'almacenamiento y un máximo de 97 conexiones simultáneas; (c) no se garantiza disponibilidad '
     'del 100 % ni SLA; (d) las conexiones SSL son obligatorias en producción.'),
    ('Integración Exclusiva con Azure Active Directory',
     'El mecanismo de inicio de sesión único (SSO) debe integrarse exclusivamente con Microsoft '
     'Azure Active Directory (Microsoft Entra ID) utilizando OAuth 2.0 y OpenID Connect. No se '
     'contempla la integración con otros proveedores de identidad (Google, Facebook, GitHub) en '
     'esta fase del proyecto. La aplicación de Azure debe estar registrada con el client ID '
     '89262870-12e6-41ab-b212-07f34b9bde0a y la authority correspondiente al tenant institucional.'),
    ('Arquitectura de API RESTful sin Estado (Stateless)',
     'El backend debe operar sin estado de sesión en el servidor. Toda la información de '
     'autenticación debe contenerse en el token JWT. No está permitido el uso de sesiones basadas '
     'en cookies ni almacenamiento en memoria del servidor. Esto garantiza que el backend pueda '
     'escalar horizontalmente sin necesidad de sticky sessions.'),
]
for titulo, desc in restricciones:
    doc.add_heading(titulo, level=3)
    add_paragraph_rich(doc, desc)

# ═══════════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del Documento —')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x7c, 0x7c, 0x8c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Clubes UNID — SRS v1.0')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x7c, 0x7c, 0x8c)

# ═══════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════
output_path = 'C:\\Users\\Equipo 1\\OneDrive\\Desktop\\SPA-Angel\\SRS_Clubes_UNID_v1.docx'
doc.save(output_path)
print(f'Documento generado exitosamente: {output_path}')
